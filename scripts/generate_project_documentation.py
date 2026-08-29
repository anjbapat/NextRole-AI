from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUT = Path("docs/NextRole-AI_Project_Documentation-v2.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)
doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)
styles["Title"].font.name = "Aptos Display"
styles["Title"].font.size = Pt(26)

title = doc.add_heading("NextRole-AI", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph("Project Documentation")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].bold = True
p = doc.add_paragraph("Stateful AI Job Application and Interview Copilot")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()


def heading(text, level=1):
    doc.add_heading(text, level=level)


def para(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix) :])
    else:
        p.add_run(text)
    return p


def bullets(items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Shading Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    return t


heading("1. Project Intention")
para(
    "NextRole-AI helps job seekers make evidence-based application decisions and prepare stronger, truthful application materials. A user uploads a resume and pastes a job description. The system extracts structured information, maps every job requirement to resume evidence, calculates an explainable fit score, recommends APPLY, REVIEW, or SKIP, prepares an application package, pauses for human approval, and generates interview preparation based on the candidate's gaps."
)
para(
    "The project is intentionally designed as an AI-engineering portfolio project rather than a general job board. Its purpose is to demonstrate structured model outputs, stateful orchestration, deterministic decision logic, human-in-the-loop control, persistence, safeguards, failure handling, and evaluation in one understandable vertical workflow."
)
heading("Primary objectives", 2)
bullets(
    [
        "Help candidates identify roles that are genuinely aligned with their evidence.",
        "Explain each recommendation instead of presenting an unsupported model-generated percentage.",
        "Tailor wording without inventing employers, technologies, metrics, responsibilities, or certifications.",
        "Keep the candidate in control before any application is considered ready.",
        "Connect missing requirements directly to interview questions and study actions.",
        "Persist the journey so applications can be reviewed and continued later.",
    ]
)

heading("2. User Journey")
numbered(
    [
        "Upload a text-based PDF resume.",
        "Paste the complete job description.",
        "Extract resume facts and job requirements into validated schemas.",
        "Map required and preferred skills to resume evidence.",
        "Calculate a deterministic weighted fit score.",
        "Route the role to APPLY, REVIEW, or SKIP.",
        "For APPLY or REVIEW roles, create an editable application package.",
        "Pause the LangGraph workflow for explicit human approval or rejection.",
        "After approval, generate interview questions, knowledge gaps, and a study plan.",
        "Save the application stage and LangGraph checkpoint in SQLite.",
    ]
)

heading("3. System Architecture")
para(
    "NextRole-AI uses one LangGraph workflow with specialized nodes. It does not create a collection of autonomous agents. This keeps behavior observable and easier to test while still demonstrating agentic state, routing, interruption, and recovery."
)
para(
    "Streamlit UI -> LangGraph workflow -> structured LLM calls + deterministic tools -> SQLite persistence"
)
doc.add_picture("docs/assets/nextrole-ai-process-handdrawn.png", width=Inches(7.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
caption = doc.add_paragraph("Figure 1. Hand-drawn overview of the complete NextRole-AI workflow.")
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
heading("Responsibility boundary", 2)
table(
    ["Large language model", "Deterministic Python"],
    [
        ["Resume and job information extraction", "Weighted fit-score calculation"],
        ["Semantic evidence comparison", "APPLY / REVIEW / SKIP thresholds"],
        ["Truthful wording suggestions", "Skill normalization and deduplication"],
        ["Recruiter message generation", "Evidence-table reconciliation"],
        ["Interview question generation", "Conditional graph routing"],
        ["Study-plan generation", "Approval enforcement and persistence"],
    ],
)

heading("4. Models and Providers")
table(
    ["Model / mode", "Role", "Status"],
    [
        [
            "Fireworks AI: accounts/fireworks/models/deepseek-v4-flash-0731",
            "Default hosted model for all structured extraction and generation tasks.",
            "Active default; verified with long-resume ResumeProfile output. Reasoning is disabled to reserve tokens for complete JSON.",
        ],
        [
            "Nebius: Qwen/Qwen3-235B-A22B",
            "Optional OpenAI-compatible alternative configured through environment variables.",
            "Supported alternative; requires NEBIUS_API_KEY.",
        ],
        [
            "Local demo mode",
            "Deterministic lexical comparison used to exercise the UI and graph without an API key.",
            "Testing/demo only; not a semantic assessment.",
        ],
    ],
)
para(
    "The application uses the OpenAI-compatible Python client for both hosted providers. Provider responses are constrained with JSON Schema and then validated with Pydantic. Empty, invalid, or truncated structured responses are retried once. The second attempt receives a larger output budget."
)
heading("Environment configuration", 2)
table(
    ["Variable", "Purpose"],
    [
        ["LLM_PROVIDER", "Selects fireworks or nebius."],
        ["FIREWORKS_API_KEY", "Fireworks authentication key."],
        ["FIREWORKS_MODEL", "Configurable Fireworks model ID."],
        ["FIREWORKS_BASE_URL", "Fireworks OpenAI-compatible endpoint."],
        ["NEBIUS_API_KEY", "Optional Nebius authentication key."],
        ["NEBIUS_MODEL", "Configurable Nebius model ID."],
        ["NEBIUS_BASE_URL", "Nebius OpenAI-compatible endpoint."],
    ],
)

heading("5. Agents in the Pipeline")
para(
    "The following names describe logical agent responsibilities. In code, they are specialized nodes within a single stateful graph, not separate autonomous processes."
)
table(
    ["Logical agent / graph node", "Purpose", "Output"],
    [
        [
            "Resume Analysis Agent / parse_resume",
            "Extracts supported candidate facts from the uploaded resume.",
            "ResumeProfile: name, summary, skills, highlights, domains, education.",
        ],
        [
            "Job Analysis Agent / analyze_job",
            "Extracts exhaustive and normalized job requirements.",
            "JobAnalysis: company, title, required/preferred skills, responsibilities, domains, education.",
        ],
        [
            "Job Fit Agent / calculate_fit",
            "Compares requirements with evidence. Python recalculates the final score and recommendation.",
            "FitAnalysis and requirement-to-evidence map.",
        ],
        [
            "Application Agent / tailor_application",
            "Reframes only existing experience for the target role.",
            "Resume suggestions, recruiter message, fit narrative, keywords, cover-letter summary.",
        ],
        [
            "Human Approval Gate / human_approval",
            "Interrupts execution until the user approves or rejects the editable package.",
            "Explicit approval decision.",
        ],
        [
            "Interview Agent / interview_prep",
            "Builds questions and preparation actions from role requirements and gaps.",
            "Categorized questions, gap priorities, preparation hints, study plan.",
        ],
        [
            "Persistence nodes / save_*",
            "Records skipped, rejected, or ready-to-submit outcomes.",
            "SQLite application record and durable graph checkpoint.",
        ],
    ],
)

heading("6. Deterministic Tools and Capabilities")
table(
    ["Tool / capability", "Responsibility"],
    [
        ["PDF text extraction", "Reads selectable text from uploaded PDF resumes with PyPDF."],
        [
            "StructuredLLM.structured",
            "Calls Fireworks or Nebius, requests JSON Schema output, validates it, detects truncation, retries with a larger budget, and normalizes provider errors.",
        ],
        [
            "Skill normalization",
            "Deduplicates skill names and removes generic orphan tokens such as 'big' or 'analysis' while preserving phrases such as 'Big Data'.",
        ],
        [
            "Evidence reconciliation",
            "Forces one evidence row per actual required/preferred job skill and removes unrelated model-generated gaps.",
        ],
        [
            "Hybrid scoring engine",
            "Calculates the final score from five components using fixed weights.",
        ],
        ["Fit router", "Routes scores to APPLY, REVIEW, or SKIP."],
        [
            "Approval router",
            "Sends approved packages to interview preparation and rejected packages to persistence.",
        ],
        [
            "SQLite application store",
            "Stores application summaries, status, stage, and timestamps.",
        ],
        [
            "SQLite LangGraph checkpointer",
            "Persists thread state and interrupts so a saved journey can be reopened.",
        ],
        [
            "Streamlit tracker",
            "Lists saved applications and changes the active thread when Continue is selected.",
        ],
    ],
)

heading("7. Fit-Scoring Method")
para(
    "The language model extracts evidence and supplies bounded component assessments. Deterministic Python calculates the overall score:"
)
table(
    ["Component", "Weight"],
    [
        ["Skills match", "40%"],
        ["Experience match", "30%"],
        ["Domain match", "15%"],
        ["Education match", "5%"],
        ["Preference match", "10%"],
    ],
)
table(
    ["Score range", "Recommendation"],
    [
        ["75-100", "APPLY"],
        ["55-74", "REVIEW"],
        ["0-54", "SKIP"],
    ],
)
para(
    "This design keeps routing reproducible. The model cannot change the thresholds or directly control the final decision."
)

heading("8. Human-in-the-Loop and Safety")
bullets(
    [
        "The application package is editable before approval.",
        "The graph visibly pauses at a LangGraph interrupt.",
        "The system cannot submit an application, fill an external form, send an email, or contact a recruiter.",
        "Approval means READY_TO_SUBMIT; the user still submits externally.",
        "Unsupported job requirements remain gaps.",
        "Tailoring instructions explicitly prohibit invented employers, skills, technologies, metrics, responsibilities, and certifications.",
        "Raw resume text is excluded from the application tracker record. It remains in the LangGraph checkpoint only so interrupted work can resume.",
    ]
)

heading("9. Current Version")
bullets(
    [
        "Streamlit interface with Job Analysis, Application, Interview Prep, and Applications tabs.",
        "PDF resume upload and pasted job-description input.",
        "Fireworks DeepSeek V4 Flash default provider with reasoning disabled, plus optional Nebius.",
        "Local fallbacks across resume, job, fit, and interview stages.",
        "Pydantic schemas for every major model response.",
        "Complete required/preferred skill extraction with deterministic cleanup.",
        "Requirement-to-resume evidence table.",
        "Deterministic hybrid scoring and conditional graph routing.",
        "Truthful, editable application package.",
        "Visible approve/reject interrupt.",
        "Gap-connected interview questions with category, difficulty, rationale, and preparation hints.",
        "Gap-to-action study plan.",
        "SQLite application tracker and durable LangGraph checkpoints.",
        "No-key demo mode.",
        "Automated tests for scoring, routing, persistence, approval, rejection, skill filtering, and evidence completeness.",
    ]
)

heading("10. Current Limitations")
bullets(
    [
        "Only text-based PDFs are supported; scanned PDFs require OCR before upload.",
        "Preference scoring currently defaults to a neutral/full value because the UI does not yet collect detailed candidate preferences.",
        "Component assessments are model-assisted even though the final weighted calculation is deterministic.",
        "Demo mode uses lexical matching and should not be treated as a reliable semantic evaluation.",
        "There is no job-search API, company research, or multi-job comparison.",
        "There is no external application submission or communication automation.",
        "SQLite is appropriate for a local portfolio MVP but not a multi-user hosted deployment without additional design.",
    ]
)

heading("11. Next Version Roadmap")
heading("Version 2 - Better matching and tracking", 2)
bullets(
    [
        "Collect candidate preferences: location, remote/hybrid/onsite, salary range, sponsorship, industry, and role level.",
        "Calculate preference score from explicit rules instead of a default value.",
        "Add OCR support for scanned resumes.",
        "Allow DOCX and plain-text resumes.",
        "Compare multiple jobs side by side.",
        "Add a mocked sample-job search abstraction, followed later by an authorized job API.",
        "Add application-stage editing: APPLIED, INTERVIEWING, REJECTED, OFFER.",
        "Add downloadable tailored package and interview-prep exports.",
        "Add evaluation cases measuring extraction completeness, evidence correctness, hallucination rate, and routing stability.",
    ]
)
heading("Version 3 - Integrations and memory", 2)
bullets(
    [
        "Authorized Gmail integration for draft-only recruiter follow-ups.",
        "Calendar integration for interview preparation and reminders.",
        "Optional Mem0-based long-term preference memory with user controls.",
        "Optional Pinecone semantic retrieval for multiple resume versions, portfolio projects, and supporting evidence.",
        "Company research with cited sources.",
        "Notifications for follow-up dates and interview milestones.",
    ]
)
heading("Version 4 - Production evolution", 2)
bullets(
    [
        "Multi-user authentication, access controls, encryption, retention controls, and audit logging.",
        "PostgreSQL-backed persistence and hosted checkpointing.",
        "Observability and model-quality monitoring.",
        "Human feedback loops that recalibrate scoring without allowing opaque automatic decisions.",
        "Carefully scoped multi-agent decomposition only where independent parallel work provides measurable value.",
        "Authorized browser or application integrations subject to platform rules and explicit confirmation.",
    ]
)

heading("12. Technology Stack")
table(
    ["Layer", "Technology"],
    [
        ["User interface", "Streamlit"],
        ["Workflow orchestration", "LangGraph"],
        ["Schemas and validation", "Pydantic"],
        ["Hosted inference", "Fireworks AI; optional Nebius"],
        ["Provider client", "OpenAI-compatible Python SDK"],
        ["PDF parsing", "PyPDF"],
        ["State and tracker persistence", "SQLite and LangGraph SQLite checkpointer"],
        ["Configuration", "python-dotenv"],
        ["Testing and quality", "pytest and Ruff"],
        ["Dependency management", "uv"],
    ],
)

heading("13. Portfolio Summary")
para(
    "NextRole-AI demonstrates where an LLM adds value and where deterministic software should remain in control. Models perform semantic extraction, evidence comparison, rewriting, and question generation. Python owns validation, normalization, scoring, routing, permissions, approval, and persistence. The centerpiece is the traceable connection: job requirement -> resume evidence -> identified gap -> interview question -> study action."
)

doc.add_page_break()
heading("Appendix: Workflow State", 1)
bullets(
    [
        "journey_id and job description",
        "structured resume profile and job analysis",
        "component scores, overall score, and recommendation",
        "requirement-to-evidence mapping",
        "application package and approval status",
        "interview questions, knowledge gaps, and study plan",
        "current workflow step, status, and error collection",
    ]
)
para("End of document.")
doc.save(OUT)
print(OUT.resolve())
